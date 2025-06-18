# logging messages to monitor execution flow.
import os
import re
import logging
from typing import Optional, List, Dict, Tuple
from pypdf import PdfReader
from docx import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('file_processor.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx']
        logger.info("Initialized FileProcessor supporting: %s", 
                   ", ".join(self.supported_formats))

    def _log_processing_start(self, file_path: str, file_type: str):
        logger.info("┌[START] Processing %s file: %s", 
                   file_type.upper(), os.path.basename(file_path))

    def _log_processing_end(self, file_path: str, success: bool, pages: int = 0):
        status = "SUCCESS" if success else "FAILED"
        log_msg = f"└[{status}] Processed {os.path.basename(file_path)}"
        if pages > 0:
            log_msg += f" ({pages} pages)"
        logger.info(log_msg)

    def extract_text_from_pdf(self, file_path: str) -> Optional[str]:
        self._log_processing_start(file_path, "PDF")
        try:
            reader = PdfReader(file_path)
            text = ""
            for i, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += re.sub(r'\s+', ' ', page_text).strip() + "\n"
                if i % 10 == 0 or i == len(reader.pages):
                    logger.debug("  ├ Processed page %d/%d", i, len(reader.pages))
            
            self._log_processing_end(file_path, True, len(reader.pages))
            return text.strip()
        except Exception as e:
            logger.error("  ├ Error: %s", str(e))
            self._log_processing_end(file_path, False)
            return None

    def extract_text_from_txt(self, file_path: str) -> Optional[str]:
        self._log_processing_start(file_path, "TXT")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                char_count = len(content)
                self._log_processing_end(file_path, True)
                logger.debug("  ├ Extracted %d characters", char_count)
                return content
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                    self._log_processing_end(file_path, True)
                    logger.warning("  ├ Used latin-1 fallback encoding")
                    return content
            except Exception as e:
                logger.error("  ├ Error: %s", str(e))
                self._log_processing_end(file_path, False)
                return None

    def extract_text_from_docx(self, file_path: str) -> Optional[str]:
        self._log_processing_start(file_path, "DOCX")
        try:
            doc = Document(file_path)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            self._log_processing_end(file_path, True)
            logger.debug("  ├ Extracted %d paragraphs", len(paragraphs))
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error("  ├ Error: %s", str(e))
            self._log_processing_end(file_path, False)
            return None

    def process_file(self, file_path: str) -> Optional[str]:
        if not os.path.exists(file_path):
            logger.error("File not found: %s", file_path)
            return None

        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.supported_formats:
            logger.error("Unsupported format: %s", file_ext)
            return None

        processors = {
            '.pdf': self.extract_text_from_pdf,
            '.txt': self.extract_text_from_txt,
            '.docx': self.extract_text_from_docx
        }
        return processors[file_ext](file_path)

    def batch_process(self, directory: str) -> Dict[str, str]:
        logger.info("=== Starting batch processing of %s ===", directory)
        results = {}
        for filename in os.listdir(directory):
            file_path = os.path.join(directory, filename)
            if os.path.isfile(file_path):
                content = self.process_file(file_path)
                if content:
                    results[filename] = content
        logger.info("=== Batch complete: %d/%d files processed ===", 
                   len(results), len(os.listdir(directory)))
        return results

# from pypdf import PdfReader
# from docx import Document
# import os
# import logging
# from typing import Optional, Union
# import re

# # Configure logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# class FileProcessor:
#     def __init__(self):
#         self.supported_formats = ['.pdf', '.txt', '.docx']

#     def extract_text_from_pdf(self, file_path: str) -> str:
#         """Extract text from PDF with improved error handling and formatting"""
#         try:
#             reader = PdfReader(file_path)
#             text = ""
#             for page in reader.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     # Clean up PDF text artifacts
#                     text += re.sub(r'\s+', ' ', page_text).strip() + "\n"
#             return text.strip()
#         except Exception as e:
#             logger.error(f"Error reading PDF {file_path}: {str(e)}")
#             raise

#     def extract_text_from_txt(self, file_path: str) -> str:
#         """Read text file with encoding fallback"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as f:
#                 return f.read()
#         except UnicodeDecodeError:
#             try:
#                 with open(file_path, 'r', encoding='latin-1') as f:
#                     return f.read()
#             except Exception as e:
#                 logger.error(f"Error reading TXT {file_path}: {str(e)}")
#                 raise

#     def extract_text_from_docx(self, file_path: str) -> str:
#         """Extract text from DOCX with structure preservation"""
#         try:
#             doc = Document(file_path)
#             return "\n".join(
#                 para.text.strip() 
#                 for para in doc.paragraphs 
#                 if para.text.strip()
#             )
#         except Exception as e:
#             logger.error(f"Error reading DOCX {file_path}: {str(e)}")
#             raise

#     def _structure_crm_data(self, content: str) -> str:
#         """Enhanced CRM data structuring with clear sections"""
#         structured = []
#         current_section = None
        
#         for line in content.split('\n'):
#             if line.startswith('###'):
#                 current_section = line.strip('#').strip()
#                 structured.append(f"\nSECTION: {current_section}\n")
#             elif current_section and line.strip():
#                 structured.append(line)
        
#         return '\n'.join(structured)

#     def process_file(self, file_path: str) -> Optional[str]:
#         """Process any supported file type with special CRM handling"""
#         try:
#             if not os.path.exists(file_path):
#                 logger.warning(f"File not found: {file_path}")
#                 return None

#             file_ext = os.path.splitext(file_path)[1].lower()
#             if file_ext not in self.supported_formats:
#                 raise ValueError(f"Unsupported file format: {file_ext}")

#             # Special handling for CRM data file
#             if os.path.basename(file_path).lower() == 'info.txt':
#                 content = self.extract_text_from_txt(file_path)
#                 return self._structure_crm_data(content)

#             # Normal file processing
#             if file_ext == '.pdf':
#                 return self.extract_text_from_pdf(file_path)
#             elif file_ext == '.txt':
#                 return self.extract_text_from_txt(file_path)
#             elif file_ext == '.docx':
#                 return self.extract_text_from_docx(file_path)

#         except Exception as e:
#             logger.error(f"Failed to process {file_path}: {str(e)}")
#             return None

#     def batch_process(self, directory: str) -> dict:
#         """Process all supported files in a directory"""
#         results = {}
#         for filename in os.listdir(directory):
#             file_path = os.path.join(directory, filename)
#             try:
#                 if os.path.isfile(file_path):
#                     content = self.process_file(file_path)
#                     if content:
#                         results[filename] = content
#             except Exception as e:
#                 logger.error(f"Skipping {filename}: {str(e)}")
#         return results
# # from pypdf import PdfReader
# # from docx import Document
# # import os

# # def extract_text_from_pdf(file_path):
# #     reader = PdfReader(file_path)
# #     text = ""
# #     for page in reader.pages:
# #         text += page.extract_text()
# #     return text

# # def extract_text_from_txt(file_path):
# #     with open(file_path, 'r', encoding='utf-8') as f:
# #         return f.read()

# # def extract_text_from_docx(file_path):
# #     doc = Document(file_path)
# #     return "\n".join([para.text for para in doc.paragraphs])

# # def process_file(file_path):
# #     if file_path.endswith('.pdf'):
# #         return extract_text_from_pdf(file_path)
# #     elif file_path.endswith('.txt'):
# #         return extract_text_from_txt(file_path)
# #     elif file_path.endswith('.docx'):
# #         return extract_text_from_docx(file_path)
# #     else:
# #         raise ValueError(f"Unsupported file format: {file_path}")